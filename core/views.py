# core/views.py

import os
import io
import textwrap

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login
from django.http import (
    HttpResponse,
    JsonResponse,
    HttpResponseBadRequest
)

from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
import json



from django.core.paginator import Paginator
from django.template.loader import render_to_string
from django.conf import settings

from .forms import AudioUploadForm, SignUpForm
from .models import AudioFile, ImportBatch, ImportItem
from .tasks import process_audio_file, discover_link

# Word export
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# PDF export
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import arabic_reshaper
from bidi.algorithm import get_display

from celery import current_app


# ---------------------------------------------------------
# Utility Helpers
# ---------------------------------------------------------

def clean_text_for_export(text):
    """Replaces emojis with export-safe text."""
    if not text:
        return ""
    repl = {
        '🕒': ' [زمان]: ',
        '🎵': ' [موسیقی]: ',
        '🆔': ' [گوینده]: ',
        '✔': ' [تیک] ',
        '⚠': ' [هشدار] ',
    }
    for k, v in repl.items():
        text = text.replace(k, v)
    return text


def get_safe_filename(audio_file, ext):
    """Generates a safe, user-friendly filename."""
    if audio_file.audio_file:
        try:
            original = os.path.basename(audio_file.audio_file.name)
            name = os.path.splitext(original)[0]
            clean = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).strip()
            return f"{clean}.{ext}"
        except:
            pass
    return f"transcript_{audio_file.id}.{ext}"


# ---------------------------------------------------------
# Landing + Authentication
# ---------------------------------------------------------

def landing(request):
    if request.user.is_authenticated:
        return redirect("dashboard")
    return render(request, "core/landing.html")



def signup(request):
    if request.method == "POST":
        form = SignUpForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect("dashboard")  # ✅ مهم
    else:
        form = SignUpForm()

    return render(request, "registration/signup.html", {"form": form})


# ---------------------------------------------------------
# Dashboard (GET)
# Table updates via AJAX (polling)
# ---------------------------------------------------------

@login_required
def dashboard(request):
    """Main dashboard page."""
    files = AudioFile.objects.filter(user=request.user).order_by("-created_at")
    paginator = Paginator(files, 10)
    page = request.GET.get("page")
    page_obj = paginator.get_page(page)

    return render(request, "core/dashboard.html", {"files": page_obj})


# ---------------------------------------------------------
# Upload (AJAX)
# ---------------------------------------------------------


@login_required
def upload_file(request):
    """
    AJAX Upload (Refresh-safe)

    Flow:
    1) Create DB row immediately (status=UPLOADING)
    2) Bind form to the same instance
    3) If upload interrupted → mark FAILED with human message
    4) If success → move to PENDING / PROCESSING queue
    """

    if request.method != "POST":
        return HttpResponseBadRequest("Invalid request")

    # -------------------------------------------------
    # 1. Create DB row FIRST (critical for refresh)
    # -------------------------------------------------
    audio = AudioFile.objects.create(
        user=request.user,
        status=AudioFile.Status.UPLOADING
    )

    # -------------------------------------------------
    # 2. Bind form to EXISTING instance
    # -------------------------------------------------
    form = AudioUploadForm(
        request.POST,
        request.FILES,
        instance=audio,
        user=request.user
    )

    # -------------------------------------------------
    # 3. Validation failed (refresh / cancel / broken upload)
    # -------------------------------------------------
    if not form.is_valid():
        audio.status = AudioFile.Status.FAILED
        audio.error_message = (
            "آپلود ناقص بود. "
            "احتمالاً صفحه رفرش شده یا ارتباط قطع شده است."
        )
        audio.save(update_fields=["status", "error_message"])

        return JsonResponse(
            {
                "success": False,
                "message": audio.error_message
            },
            status=400
        )

    # -------------------------------------------------
    # 4. Save uploaded file on SAME row
    # -------------------------------------------------
    audio = form.save(commit=False)

    # -------------------------------------------------
    # 5. Detect VIDEO file
    # -------------------------------------------------
    filename = (audio.audio_file.name or "").lower()
    video_exts = (".mp4", ".mkv", ".avi", ".mov", ".webm")
    audio.is_video = filename.endswith(video_exts)

    audio.save()

    # -------------------------------------------------
    # 6. Queue logic (one active job per user)
    # -------------------------------------------------
    user_has_active_job = AudioFile.objects.filter(
        user=request.user,
        status__in=[
            AudioFile.Status.PENDING,
            AudioFile.Status.PROCESSING
        ]
    ).exclude(id=audio.id).exists()

    audio.status = AudioFile.Status.PENDING
    audio.save(update_fields=["status"])

    if not user_has_active_job:
        task = process_audio_file.delay(audio.id)
        audio.task_id = task.id
        audio.save(update_fields=["task_id"])

    # -------------------------------------------------
    # 7. Success response
    # -------------------------------------------------
    return JsonResponse(
        {
            "success": True,
            "file_id": audio.id
        }
    )
    
    
    
@login_required
def create_import_batch(request):
    """
    Creates an ImportBatch from a URL and starts discover task.
    """
    if request.method != "POST":
        return HttpResponseBadRequest("Invalid request")

    url = request.POST.get("url", "").strip()
    if not url:
        return JsonResponse({
            "success": False,
            "error": "لینک معتبر وارد نشده است."
        }, status=400)

    batch = ImportBatch.objects.create(
        user=request.user,
        source_url=url,
        status=ImportBatch.Status.CREATED
    )
    discover_link.delay(batch.id)

    # Start discover task
    # task = discover_link.delay(batch.id)
    # batch.task_id = task.id
    # batch.save(update_fields=["task_id"])

    return JsonResponse({
        "success": True,
        "batch_id": batch.id
    })


@login_required
def import_batch_status(request, batch_id):
    batch = get_object_or_404(
        ImportBatch,
        id=batch_id,
        user=request.user
    )

    if batch.status == ImportBatch.Status.READY:
        items = batch.items.all()
        html = render_to_string(
            "partials/import_items.html",
            {"batch": batch, "items": items},
            request=request
        )

        return JsonResponse({
            "status": "READY",
            "html": html
        })

    if batch.status == ImportBatch.Status.FAILED:
        return JsonResponse({
            "status": "FAILED",
            "error": batch.error_message or "خطا در خواندن لینک."
        })

    return JsonResponse({
        "status": batch.status
    })
    
    
@login_required
@require_POST
def enqueue_import_items(request):
    """
    Creates AudioFile records from selected ImportItems and enqueues them
    respecting user-level queue.
    """

    # --------------------------------------------------
    # Parse request data (JSON preferred)
    # --------------------------------------------------
    try:
        data = json.loads(request.body)
        item_ids = data.get("items", [])
        batch_id = data.get("batch_id")
        model_name = data.get("model_name", "eboo")
    except Exception:
        # fallback for form-data
        item_ids = request.POST.getlist("items[]")
        batch_id = request.POST.get("batch_id")
        model_name = request.POST.get("model_name", "eboo")

    if not item_ids or not batch_id:
        return JsonResponse({
            "success": False,
            "error": "هیچ فایلی انتخاب نشده است."
        }, status=400)

    # --------------------------------------------------
    # Load batch (security: user-owned)
    # --------------------------------------------------
    batch = get_object_or_404(
        ImportBatch,
        id=batch_id,
        user=request.user
    )

    # --------------------------------------------------
    # Load items
    # --------------------------------------------------
    items = ImportItem.objects.filter(
        id__in=item_ids,
        batch=batch
    )

    if not items.exists():
        return JsonResponse({
            "success": False,
            "error": "آیتم معتبری یافت نشد."
        }, status=400)

    created_files = []

    # --------------------------------------------------
    # Create AudioFile objects
    # --------------------------------------------------
    for item in items:
        audio = AudioFile.objects.create(
            user=request.user,
            title=item.title or "Imported file",
            source_url=item.source_url,
            is_video=item.is_video,
            model_name=model_name,
            status=AudioFile.Status.PENDING,
            import_batch=batch,
        )
        created_files.append(audio)

    # --------------------------------------------------
    # User-level queue logic
    # --------------------------------------------------
    has_active_job = AudioFile.objects.filter(
        user=request.user,
        status=AudioFile.Status.PROCESSING
    ).exists()

    # If user is idle → start first job
    if created_files and not has_active_job:
        first = created_files[0]
        task = process_audio_file.delay(first.id)
        first.task_id = task.id
        first.save(update_fields=["task_id"])

    return JsonResponse({
        "success": True,
        "created": len(created_files)
    })
    
    

# ---------------------------------------------------------
# AJAX — File list refresh for polling
# ---------------------------------------------------------

@login_required
def get_files(request):
    """Returns updated table HTML for polling."""
    files = AudioFile.objects.filter(user=request.user).order_by("-created_at")
    paginator = Paginator(files, 10)
    page = request.GET.get("page")
    page_obj = paginator.get_page(page)

    html = render_to_string(
        "partials/file_table_container.html",
        {"files": page_obj},
        request=request
    )
    return HttpResponse(html)


@login_required
def update_row(request, file_id):
    """Returns HTML for a single table row."""
    file = get_object_or_404(AudioFile, id=file_id, user=request.user)
    html = render_to_string("partials/row.html", {"file": file}, request=request)
    return HttpResponse(html)


# ---------------------------------------------------------
# Delete File (AJAX + JSON response)
# With REAL Celery terminate
# ---------------------------------------------------------

@login_required
def delete_file(request, pk):
    """
    AJAX-only deletion endpoint.
    - Accepts DELETE request.
    - Kills Celery task if still running.
    - Removes DB entry.
    - Returns JSON.
    """
    if request.method != "DELETE":
        return HttpResponseBadRequest("Invalid request")

    obj = get_object_or_404(AudioFile, pk=pk, user=request.user)

    # Kill Celery task if in progress
    # if obj.task_id:
    #     try:
    #         app = current_app
    #         app.control.revoke(obj.task_id, terminate=True, signal='SIGKILL')
    #     except Exception as e:
    #         print("Celery revoke error:", e)

    obj.delete()

    return JsonResponse({"success": True})


# ---------------------------------------------------------
# Downloads
# ---------------------------------------------------------

@login_required
def download_txt(request, file_id):
    """Exports transcript as TXT."""
    audio_file = get_object_or_404(AudioFile, id=file_id, user=request.user)
    raw = audio_file.transcript_text or "متنی موجود نیست."
    txt = clean_text_for_export(raw)

    content = f"عنوان: {audio_file.title}\n------------------\n{txt}"

    response = HttpResponse(content, content_type="text/plain; charset=utf-8")
    filename = get_safe_filename(audio_file, "txt")
    response['Content-Disposition'] = f"attachment; filename*=UTF-8''{filename}"
    return response


@login_required
def download_word(request, file_id):
    """Exports transcript as a Word DOCX file."""
    audio_file = get_object_or_404(AudioFile, id=file_id, user=request.user)

    doc = Document()
    title_p = doc.add_heading(audio_file.title or "متن استخراج شده", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    raw = audio_file.transcript_text or "متنی موجود نیست."
    clean = clean_text_for_export(raw)

    for line in clean.split("\n"):
        if line.strip():
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.bidi = True
            if p.runs:
                r = p.runs[0]
                r.font.name = "Arial"
                r.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    filename = get_safe_filename(audio_file, "docx")
    response['Content-Disposition'] = f"attachment; filename*=UTF-8''{filename}"
    return response


@login_required
def download_pdf(request, file_id):
    """Exports transcript as PDF."""
    audio_file = get_object_or_404(AudioFile, id=file_id, user=request.user)

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    margin_x, margin_y = 50, 50

    font_name = "Helvetica"
    try:
        path = os.path.join(settings.BASE_DIR, "core", "static", "fonts", "Vazir.ttf")
        pdfmetrics.registerFont(TTFont("Vazir", path))
        font_name = "Vazir"
    except Exception as e:
        print("Font error:", e)

    c.setFont(font_name, 12)

    raw = audio_file.transcript_text or "متنی موجود نیست."
    clean = clean_text_for_export(raw)

    y = height - margin_y
    line_h = 20

    for paragraph in clean.split("\n"):
        if not paragraph.strip():
            y -= line_h
            continue

        lines = textwrap.wrap(paragraph, width=90)

        for l in lines:
            if font_name == "Vazir":
                reshaped = arabic_reshaper.reshape(l)
                bidi = get_display(reshaped)
            else:
                bidi = l

            c.drawRightString(width - margin_x, y, bidi)
            y -= line_h

            if y < margin_y:
                c.showPage()
                c.setFont(font_name, 12)
                y = height - margin_y

    c.showPage()
    c.save()

    buf.seek(0)
    response = HttpResponse(buf, content_type="application/pdf")
    filename = get_safe_filename(audio_file, "pdf")
    response['Content-Disposition'] = f"attachment; filename*=UTF-8''{filename}"
    return response
